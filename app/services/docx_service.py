"""
DocxService - Generates DOCX documents from drafts with annexes.
"""

import io
import re
from datetime import datetime
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from loguru import logger

from app.models.schemas import CaseContext, ProcessedAnnex


class DocxService:
    """
    Service for generating properly formatted DOCX documents.
    """

    # Document type to title mapping
    DOCUMENT_TITLES = {
        "demanda_ods": "DEMANDA DE OBLIGACIÓN DE DAR SUMA DE DINERO",
        "demanda_eg": "DEMANDA DE EJECUCIÓN DE GARANTÍAS",
        "demanda_leasing": "DEMANDA POR INCUMPLIMIENTO DE LEASING",
        "solicitud_remate": "SOLICITUD DE REMATE",
        "solicitud_adjudicacion": "SOLICITUD DE ADJUDICACIÓN",
        "solicitud_lanzamiento": "SOLICITUD DE LANZAMIENTO",
        "solicitud_tasacion": "SOLICITUD DE TASACIÓN",
        "recurso_apelacion": "RECURSO DE APELACIÓN",
        "recurso_casacion": "RECURSO DE CASACIÓN",
        "escrito_impulso": "ESCRITO DE IMPULSO PROCESAL",
        "escrito_subsanacion": "ESCRITO DE SUBSANACIÓN",
        "medida_cautelar_fuera": "MEDIDA CAUTELAR FUERA DE PROCESO",
        "medida_cautelar_dentro": "MEDIDA CAUTELAR DENTRO DE PROCESO",
    }

    def __init__(self):
        pass

    async def generate(
        self,
        draft: str,
        document_type: str,
        context: CaseContext,
    ) -> bytes:
        """
        Generate a DOCX document from the draft text.

        Args:
            draft: The document draft text
            document_type: Type of document
            context: Case file context

        Returns:
            DOCX file as bytes
        """
        logger.info(f"Generating DOCX for: {document_type}")

        # Create document
        doc = Document()

        # Set up styles
        self._setup_styles(doc)

        # Set page margins
        self._set_margins(doc)

        # Add header with logo placeholder
        self._add_header(doc, context)

        # Parse and add content
        self._add_content(doc, draft, document_type)

        # Add footer
        self._add_footer(doc)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    def _setup_styles(self, doc: Document) -> None:
        """Set up document styles."""
        styles = doc.styles

        # Title style
        if "DocTitle" not in [s.name for s in styles]:
            title_style = styles.add_style("DocTitle", WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = "Arial"
            title_style.font.size = Pt(14)
            title_style.font.bold = True

        # Heading style
        if "DocHeading" not in [s.name for s in styles]:
            heading_style = styles.add_style("DocHeading", WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.name = "Arial"
            heading_style.font.size = Pt(12)
            heading_style.font.bold = True

        # Body style
        if "DocBody" not in [s.name for s in styles]:
            body_style = styles.add_style("DocBody", WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = "Arial"
            body_style.font.size = Pt(11)

    def _set_margins(self, doc: Document) -> None:
        """Set page margins."""
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2.5)

    def _add_header(self, doc: Document, context: CaseContext) -> None:
        """Add document header."""
        # Add expediente info at top right
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if context.case_number and context.case_number != "NUEVO":
            header_para.add_run(f"Expediente: {context.case_number}\n").bold = True
        header_para.add_run(f"Cuaderno: Principal\n")

        # Add spacing
        doc.add_paragraph()

    def _add_content(self, doc: Document, draft: str, document_type: str) -> None:
        """Parse draft text and add formatted content with markdown support."""
        # Split content into sections
        lines = draft.split("\n")

        for line in lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue

            # Check for markdown headers
            h1_match = re.match(r"^#\s+(.+)$", stripped)
            h2_match = re.match(r"^##\s+(.+)$", stripped)
            h3_match = re.match(r"^###\s+(.+)$", stripped)
            h4_match = re.match(r"^####\s+(.+)$", stripped)

            if h1_match:
                # Main title - centered, bold, uppercase
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                text = self._strip_markdown_formatting(h1_match.group(1))
                run = para.add_run(text.upper())
                run.bold = True
                run.font.size = Pt(14)
                run.font.name = "Arial"
                continue

            if h2_match:
                # Section header - bold, uppercase
                para = doc.add_paragraph()
                text = self._strip_markdown_formatting(h2_match.group(1))
                run = para.add_run(text.upper())
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = "Arial"
                para.paragraph_format.space_before = Pt(12)
                continue

            if h3_match:
                # Subsection header - bold
                para = doc.add_paragraph()
                text = self._strip_markdown_formatting(h3_match.group(1))
                run = para.add_run(text)
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = "Arial"
                para.paragraph_format.space_before = Pt(10)
                continue

            if h4_match:
                # Minor header - bold
                para = doc.add_paragraph()
                text = self._strip_markdown_formatting(h4_match.group(1))
                run = para.add_run(text)
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = "Arial"
                continue

            # Check for horizontal rule
            if re.match(r"^[-*_]{3,}$", stripped):
                # Add horizontal line
                para = doc.add_paragraph()
                para.add_run("_" * 60)
                continue

            # Detect legal section headers (SUMILLA:, PETITORIO:, etc.)
            if self._is_section_header(stripped):
                para = doc.add_paragraph()
                # Parse inline markdown and add with formatting
                self._add_formatted_runs(para, stripped, is_header=True)
                para.paragraph_format.space_before = Pt(12)
                continue

            # Detect numbered items (1., 2., a), b), etc.)
            numbered_match = re.match(r"^(\d+[\.\)])\s+(.+)$", stripped)
            letter_match = re.match(r"^([a-z][\.\)])\s+(.+)$", stripped, re.IGNORECASE)

            if numbered_match:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Cm(1)
                # Add number
                run = para.add_run(numbered_match.group(1) + " ")
                run.font.name = "Arial"
                run.font.size = Pt(11)
                # Add rest with formatting
                self._add_formatted_runs(para, numbered_match.group(2))
                continue

            if letter_match:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Cm(1)
                run = para.add_run(letter_match.group(1) + " ")
                run.font.name = "Arial"
                run.font.size = Pt(11)
                self._add_formatted_runs(para, letter_match.group(2))
                continue

            # Detect bullet points
            bullet_match = re.match(r"^[-•*]\s+(.+)$", stripped)
            if bullet_match:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Cm(1)
                run = para.add_run("• ")
                run.font.name = "Arial"
                self._add_formatted_runs(para, bullet_match.group(1))
                continue

            # Regular text - parse inline markdown
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Cm(1.25)
            self._add_formatted_runs(para, stripped)

    def _add_formatted_runs(self, para, text: str, is_header: bool = False) -> None:
        """
        Parse inline markdown and add formatted runs to paragraph.
        Handles **bold**, *italic*, __bold__, _italic_ and combinations.
        """
        # Pattern to match markdown formatting
        # This handles: **bold**, *italic*, __bold__, _italic_
        pattern = r"(\*\*[^*]+\*\*|__[^_]+__|(?<!\*)\*[^*]+\*(?!\*)|(?<!_)_[^_]+_(?!_))"

        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue

            # Check for bold (**text** or __text__)
            bold_match = re.match(r"^\*\*(.+)\*\*$", part) or re.match(r"^__(.+)__$", part)
            if bold_match:
                run = para.add_run(bold_match.group(1))
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(12) if is_header else Pt(11)
                continue

            # Check for italic (*text* or _text_)
            italic_match = re.match(r"^\*([^*]+)\*$", part) or re.match(r"^_([^_]+)_$", part)
            if italic_match:
                run = para.add_run(italic_match.group(1))
                run.italic = True
                run.font.name = "Arial"
                run.font.size = Pt(12) if is_header else Pt(11)
                continue

            # Regular text
            run = para.add_run(part)
            run.font.name = "Arial"
            run.font.size = Pt(12) if is_header else Pt(11)
            if is_header:
                run.bold = True

    def _strip_markdown_formatting(self, text: str) -> str:
        """Remove markdown formatting symbols from text."""
        # Remove bold markers
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"__([^_]+)__", r"\1", text)
        # Remove italic markers
        text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", text)
        text = re.sub(r"(?<!_)_([^_]+)_(?!_)", r"\1", text)
        return text

    def _is_section_header(self, line: str) -> bool:
        """Detect if line is a section header."""
        headers = [
            "SUMILLA",
            "SECRETARIO",
            "EXPEDIENTE",
            "CUADERNO",
            "ESCRITO",
            "DEMANDANTE",
            "DEMANDADO",
            "MATERIA",
            "VÍA PROCEDIMENTAL",
            "VIA PROCEDIMENTAL",
            "CUANTÍA",
            "CUANTIA",
            "PETITORIO",
            "FUNDAMENTOS DE HECHO",
            "FUNDAMENTOS DE DERECHO",
            "MEDIOS PROBATORIOS",
            "ANEXOS",
            "OTROSI",
            "OTROSÍ",
            "PRIMER OTROSÍ",
            "SEGUNDO OTROSÍ",
            "I.",
            "II.",
            "III.",
            "IV.",
            "V.",
            "VI.",
            "VII.",
            "VIII.",
            "IX.",
            "X.",
        ]

        line_upper = line.upper().strip()

        for header in headers:
            if line_upper.startswith(header):
                return True

        # Check for numbered roman sections
        if re.match(r"^[IVX]+\.\s", line_upper):
            return True

        return False

    def _add_footer(self, doc: Document) -> None:
        """Add document footer with signature area."""
        doc.add_paragraph()
        doc.add_paragraph()

        # City and date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        today = datetime.now().strftime("%d de %B de %Y")
        date_para.add_run(f"Lima, {today}")

        doc.add_paragraph()
        doc.add_paragraph()

        # Signature line
        sig_para = doc.add_paragraph()
        sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sig_para.add_run("_" * 40)

        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.add_run("[NOMBRE DEL ABOGADO]")

        reg_para = doc.add_paragraph()
        reg_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        reg_para.add_run("Reg. CAL N° [NÚMERO]")

    async def generate_from_template(
        self,
        template_name: str,
        context: CaseContext,
        variables: dict,
    ) -> bytes:
        """
        Generate a document from a template file.

        Args:
            template_name: Name of the template file
            context: Case file context
            variables: Variables to replace in template

        Returns:
            DOCX file as bytes
        """
        # TODO: Implement template-based generation
        raise NotImplementedError("Template-based generation not yet implemented")

    async def generate_with_annexes(
        self,
        draft: str,
        document_type: str,
        context: CaseContext,
        annexes: Optional[list[ProcessedAnnex]] = None,
    ) -> bytes:
        """
        Generate a DOCX document with embedded annexes.

        Args:
            draft: The document draft text
            document_type: Type of document
            context: Case file context
            annexes: List of processed annexes with image data

        Returns:
            DOCX file as bytes
        """
        logger.info(f"Generating DOCX with {len(annexes or [])} annexes for: {document_type}")

        # Create document
        doc = Document()

        # Set up styles
        self._setup_styles(doc)

        # Set page margins
        self._set_margins(doc)

        # Add header with logo placeholder
        self._add_header(doc, context)

        # Parse and add content
        self._add_content(doc, draft, document_type)

        # Add footer
        self._add_footer(doc)

        # Add annexes section if provided
        if annexes and len(annexes) > 0:
            self._add_annexes_section(doc, annexes)

        # Add page numbers
        self._add_page_numbers(doc)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        logger.info(f"DOCX created with {len(annexes or [])} annexes")
        return buffer.getvalue()

    def _add_page_numbers(self, doc: Document) -> None:
        """Add page numbers to document footer."""
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False

            # Clear existing footer content
            for para in footer.paragraphs:
                para.clear()

            # Add page number
            para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add page number field
            run = para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')

            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            run._r.append(fldChar3)

            run.font.size = Pt(9)
            run.font.name = "Arial"

    def _add_annexes_section(self, doc: Document, annexes: list[ProcessedAnnex]) -> None:
        """
        Add annexes section with embedded images.

        Each annex includes:
        - Page break
        - Annex header (ANEXO 1, ANEXO 2, etc.)
        - Annex name
        - Embedded image(s)
        """
        logger.info(f"Adding {len(annexes)} annexes to document")

        # Add page break before annexes
        doc.add_page_break()

        # Add main ANEXOS header
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run("ANEXOS")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Arial"

        doc.add_paragraph()  # Spacing

        for i, annex in enumerate(annexes, 1):
            # Add annex header
            annex_header = doc.add_paragraph()
            annex_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = annex_header.add_run(f"ANEXO {i}")
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = "Arial"

            # Add annex name/description
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = name_para.add_run(annex.name)
            run.font.size = Pt(11)
            run.font.name = "Arial"
            run.italic = True

            # Add source label
            source_label = self._get_source_label(annex.source.value)
            source_para = doc.add_paragraph()
            source_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = source_para.add_run(f"[{source_label}]")
            run.font.size = Pt(9)
            run.font.name = "Arial"
            run.font.color.rgb = RGBColor(128, 128, 128)

            doc.add_paragraph()  # Spacing

            # Add images
            for j, img_data in enumerate(annex.image_data):
                try:
                    # Create image stream
                    img_stream = io.BytesIO(img_data)

                    # Add centered paragraph for image
                    img_para = doc.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Calculate max width (page width minus margins)
                    max_width = Inches(6)  # ~15cm, fits within margins
                    run = img_para.add_run()
                    run.add_picture(img_stream, width=max_width)

                    # Add page indicator for multi-page PDFs
                    if annex.image_count > 1:
                        page_para = doc.add_paragraph()
                        page_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = page_para.add_run(f"Página {j + 1} de {annex.image_count}")
                        run.font.size = Pt(9)
                        run.font.name = "Arial"
                        run.font.color.rgb = RGBColor(128, 128, 128)

                except Exception as e:
                    logger.error(f"Failed to add image for annex {i}: {e}")
                    # Add error placeholder
                    error_para = doc.add_paragraph()
                    error_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = error_para.add_run("[Error al cargar imagen]")
                    run.font.color.rgb = RGBColor(255, 0, 0)

            # Add page break between annexes (except last one)
            if i < len(annexes):
                doc.add_page_break()

    def _get_source_label(self, source: str) -> str:
        """Get human-readable label for annex source."""
        labels = {
            "judicial_binnacle": "Expediente Judicial",
            "judicial_collateral": "Documento de Garantía",
            "extrajudicial_client": "Archivo del Cliente",
            "extrajudicial_payment": "Comprobante de Pago",
        }
        return labels.get(source, source)

    def generate_filename(
        self,
        document_type: str,
        case_number: Optional[str] = None,
        extension: str = "docx",
    ) -> str:
        """Generate a suggested filename for the document."""
        # Get title or use document type
        title = self.DOCUMENT_TITLES.get(document_type, document_type)
        # Clean for filename
        title_clean = re.sub(r"[^\w\s-]", "", title).strip()
        title_clean = re.sub(r"\s+", "_", title_clean)

        # Add case number if available
        if case_number and case_number != "NUEVO":
            case_clean = re.sub(r"[^\w-]", "_", case_number)
            filename = f"{title_clean}_{case_clean}"
        else:
            filename = title_clean

        # Add timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{filename}_{timestamp}"

        return f"{filename}.{extension}"
