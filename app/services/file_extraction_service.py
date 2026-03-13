"""
File Extraction Service - Extracts content from various file types.

Supports:
- PDF (text + images via Claude Vision)
- Word (.docx)
- Excel (.xlsx)
- Images (via Claude Vision)
"""

import base64
import hashlib
import io
from typing import Any, Optional

from loguru import logger

# PDF extraction
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    logger.warning("PyMuPDF not installed, PDF image extraction disabled")

try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False
    logger.warning("pypdf not installed, PDF text extraction disabled")

# Word extraction
try:
    from docx import Document as WordDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed, Word extraction disabled")

# Excel extraction
try:
    from openpyxl import load_workbook
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False
    logger.warning("openpyxl not installed, Excel extraction disabled")

# Image processing
try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False
    logger.warning("Pillow not installed, image processing disabled")

# Anthropic for Vision
try:
    import anthropic
    HAS_ANTHROPIC = True
except ImportError:
    HAS_ANTHROPIC = False
    logger.warning("anthropic not installed, image analysis disabled")

from app.config import settings


class ExtractionResult:
    """Result of file extraction."""

    def __init__(
        self,
        text: str = "",
        images: list[dict] = None,
        page_count: int = 0,
        image_count: int = 0,
        tokens_estimated: int = 0,
        error: Optional[str] = None,
    ):
        self.text = text
        self.images = images or []
        self.page_count = page_count
        self.image_count = image_count
        self.tokens_estimated = tokens_estimated
        self.error = error

    @property
    def success(self) -> bool:
        return self.error is None and (self.text or self.images)


class FileExtractionService:
    """Service for extracting content from various file types."""

    def __init__(self):
        if HAS_ANTHROPIC:
            self.anthropic_client = anthropic.Anthropic(
                api_key=settings.anthropic_api_key
            )
        else:
            self.anthropic_client = None

    def get_file_hash(self, content: bytes) -> str:
        """Calculate SHA256 hash of file content."""
        return hashlib.sha256(content).hexdigest()

    def detect_file_type(self, filename: str, content: bytes) -> str:
        """Detect file type from filename and content."""
        filename_lower = filename.lower()

        if filename_lower.endswith(".pdf"):
            return "pdf"
        elif filename_lower.endswith(".docx"):
            return "docx"
        elif filename_lower.endswith((".xlsx", ".xls")):
            return "xlsx"
        elif filename_lower.endswith((".png", ".jpg", ".jpeg", ".gif", ".webp")):
            return "image"

        # Check magic bytes
        if content[:4] == b"%PDF":
            return "pdf"
        elif content[:4] == b"PK\x03\x04":  # ZIP-based (docx, xlsx)
            if b"word/" in content[:2000]:
                return "docx"
            elif b"xl/" in content[:2000]:
                return "xlsx"

        return "other"

    async def extract(
        self,
        content: bytes,
        filename: str,
        max_pages: int = 10,
        max_images: int = 5,
        analyze_images: bool = True,
    ) -> ExtractionResult:
        """
        Extract content from a file.

        Args:
            content: File content as bytes
            filename: Original filename
            max_pages: Maximum pages to extract (for PDFs)
            max_images: Maximum images to analyze
            analyze_images: Whether to analyze images with Claude Vision

        Returns:
            ExtractionResult with extracted text and image descriptions
        """
        file_type = self.detect_file_type(filename, content)
        logger.info(f"Extracting {file_type} file: {filename}")

        try:
            if file_type == "pdf":
                return await self._extract_pdf(
                    content, max_pages, max_images, analyze_images
                )
            elif file_type == "docx":
                return self._extract_word(content)
            elif file_type == "xlsx":
                return self._extract_excel(content)
            elif file_type == "image":
                return await self._extract_image(content, filename, analyze_images)
            else:
                return ExtractionResult(
                    error=f"Unsupported file type: {file_type}"
                )
        except Exception as e:
            logger.error(f"Extraction failed for {filename}: {e}")
            return ExtractionResult(error=str(e))

    async def _extract_pdf(
        self,
        content: bytes,
        max_pages: int,
        max_images: int,
        analyze_images: bool,
    ) -> ExtractionResult:
        """Extract text and images from PDF."""
        text_parts = []
        images = []
        page_count = 0
        image_count = 0
        tokens_estimated = 0

        # Extract text with pypdf
        if HAS_PYPDF:
            try:
                pdf_file = io.BytesIO(content)
                reader = PdfReader(pdf_file)
                page_count = len(reader.pages)
                pages_to_read = min(page_count, max_pages)

                for i in range(pages_to_read):
                    page = reader.pages[i]
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Página {i + 1} ---\n{page_text}")

                if page_count > max_pages:
                    text_parts.append(
                        f"\n[... {page_count - max_pages} páginas adicionales no extraídas ...]"
                    )
            except Exception as e:
                logger.error(f"PDF text extraction failed: {e}")

        # Extract and analyze images with PyMuPDF
        if HAS_PYMUPDF and analyze_images and self.anthropic_client:
            try:
                doc = fitz.open(stream=content, filetype="pdf")
                images_extracted = 0

                for page_num in range(min(len(doc), max_pages)):
                    if images_extracted >= max_images:
                        break

                    page = doc[page_num]
                    image_list = page.get_images()

                    for img_index, img in enumerate(image_list):
                        if images_extracted >= max_images:
                            break

                        try:
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            image_bytes = base_image["image"]

                            # Only analyze if image is significant (>5KB)
                            if len(image_bytes) > 5000:
                                description = await self._analyze_image_with_vision(
                                    image_bytes,
                                    f"Imagen de página {page_num + 1}"
                                )
                                if description:
                                    images.append({
                                        "pageNumber": page_num + 1,
                                        "imageIndex": img_index,
                                        "description": description,
                                    })
                                    images_extracted += 1
                                    tokens_estimated += 500  # Estimate for vision

                        except Exception as img_error:
                            logger.debug(f"Failed to extract image: {img_error}")

                doc.close()
                image_count = images_extracted

            except Exception as e:
                logger.error(f"PDF image extraction failed: {e}")

        full_text = "\n\n".join(text_parts)
        tokens_estimated += len(full_text) // 4  # Rough estimate

        return ExtractionResult(
            text=full_text,
            images=images,
            page_count=page_count,
            image_count=image_count,
            tokens_estimated=tokens_estimated,
        )

    def _extract_word(self, content: bytes) -> ExtractionResult:
        """Extract text from Word document."""
        if not HAS_DOCX:
            return ExtractionResult(error="python-docx not installed")

        try:
            doc = WordDocument(io.BytesIO(content))
            paragraphs = []

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)

            full_text = "\n\n".join(paragraphs)

            return ExtractionResult(
                text=full_text,
                page_count=1,  # Word doesn't have pages in the same sense
                tokens_estimated=len(full_text) // 4,
            )

        except Exception as e:
            return ExtractionResult(error=f"Word extraction failed: {e}")

    def _extract_excel(self, content: bytes) -> ExtractionResult:
        """Extract text from Excel spreadsheet."""
        if not HAS_OPENPYXL:
            return ExtractionResult(error="openpyxl not installed")

        try:
            wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            sheets_text = []

            for sheet_name in wb.sheetnames[:5]:  # Limit to 5 sheets
                sheet = wb[sheet_name]
                rows_text = [f"### Hoja: {sheet_name}"]

                row_count = 0
                for row in sheet.iter_rows(max_row=100, values_only=True):  # Limit rows
                    if row_count > 100:
                        rows_text.append("[... filas adicionales omitidas ...]")
                        break

                    cell_values = [
                        str(cell) if cell is not None else ""
                        for cell in row[:20]  # Limit columns
                    ]
                    if any(cell_values):
                        rows_text.append(" | ".join(cell_values))
                        row_count += 1

                if row_count > 0:
                    sheets_text.append("\n".join(rows_text))

            wb.close()
            full_text = "\n\n".join(sheets_text)

            return ExtractionResult(
                text=full_text,
                page_count=len(wb.sheetnames),
                tokens_estimated=len(full_text) // 4,
            )

        except Exception as e:
            return ExtractionResult(error=f"Excel extraction failed: {e}")

    async def _extract_image(
        self,
        content: bytes,
        filename: str,
        analyze: bool,
    ) -> ExtractionResult:
        """Extract description from image using Claude Vision."""
        if not analyze or not self.anthropic_client:
            return ExtractionResult(
                text=f"[Imagen: {filename}]",
                image_count=1,
            )

        try:
            description = await self._analyze_image_with_vision(content, filename)
            return ExtractionResult(
                text=f"[Imagen: {filename}]\n{description}" if description else "",
                images=[{"description": description}] if description else [],
                image_count=1,
                tokens_estimated=500,
            )
        except Exception as e:
            return ExtractionResult(error=f"Image analysis failed: {e}")

    async def _analyze_image_with_vision(
        self,
        image_bytes: bytes,
        context: str = "",
    ) -> Optional[str]:
        """Analyze image using Claude Vision API."""
        if not self.anthropic_client:
            return None

        try:
            # Determine media type
            if image_bytes[:8] == b'\x89PNG\r\n\x1a\n':
                media_type = "image/png"
            elif image_bytes[:2] == b'\xff\xd8':
                media_type = "image/jpeg"
            else:
                media_type = "image/jpeg"  # Default

            # Resize if too large (Claude has limits)
            if HAS_PIL and len(image_bytes) > 1_000_000:  # 1MB
                image = Image.open(io.BytesIO(image_bytes))
                image.thumbnail((1024, 1024))
                buffer = io.BytesIO()
                image.save(buffer, format="JPEG", quality=85)
                image_bytes = buffer.getvalue()
                media_type = "image/jpeg"

            image_base64 = base64.standard_b64encode(image_bytes).decode("utf-8")

            response = self.anthropic_client.messages.create(
                model=settings.claude_model,
                max_tokens=500,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": media_type,
                                    "data": image_base64,
                                },
                            },
                            {
                                "type": "text",
                                "text": f"""Analiza esta imagen de un documento legal peruano. {context}

Describe brevemente:
1. Tipo de documento (resolución, notificación, escrito, etc.)
2. Información clave visible (fechas, números, montos, nombres)
3. Estado o sello si hay (firmado, recibido, etc.)

Responde en español, máximo 200 palabras.""",
                            },
                        ],
                    }
                ],
            )

            return response.content[0].text

        except Exception as e:
            logger.error(f"Vision analysis failed: {e}")
            return None


# Singleton instance
_file_extraction_service: Optional[FileExtractionService] = None


def get_file_extraction_service() -> FileExtractionService:
    """Get or create the file extraction service instance."""
    global _file_extraction_service
    if _file_extraction_service is None:
        _file_extraction_service = FileExtractionService()
    return _file_extraction_service
